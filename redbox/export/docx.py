import tempfile
from pprint import pprint
from typing import List, Optional

import markdown
from dateutil import parser
from docx import Document
from docx.shared import Inches
from unstructured.partition.html import partition_html

from redbox import __version__ as redbox_version
from redbox.models import File, SpotlightComplete


def lookup_indentedness(raw: str, line_str_to_match: str):
    for line in raw.split("\n"):
        if line_str_to_match in line:
            # count number of spaces at start of line
            return len(line) - len(line.lstrip(" "))


def spotlight_complete_to_docx(
    spotlight_complete: SpotlightComplete,
    files: List[File],
    title: Optional[str] = None,
):
    document = Document()

    document.styles["Normal"].font.name = "Arial"
    document.styles["Title"].font.name = "Arial"
    document.styles["Heading 1"].font.name = "Arial"
    document.styles["Heading 2"].font.name = "Arial"
    document.styles["Heading 3"].font.name = "Arial"

    # Header and Footer
    section = document.sections[0]
    header = section.header
    header.paragraphs[0].text = "\tUP TO OFFICIAL SENSITIVE"
    header.paragraphs[0].style.font.name = "Arial"
    header.paragraphs[0].style.font.bold = True
    # Add page number to header (right hand side)

    footer = section.footer
    summary_datetime = parser.parse(spotlight_complete.created_datetime)
    footer.paragraphs[0].text = (
        f"This document is AI generated\t{redbox_version}\tGenerated: {summary_datetime.strftime('%d %B %Y %H:%M')}"
    )
    footer.paragraphs[0].style.font.name = "Arial"

    # Content
    if title is not None:
        document.add_heading(title, level=0)
    else:
        document.add_heading("Redbox Copilot", level=0)

    document.add_heading("Summarised Files", level=1)
    for file in files:
        document.add_paragraph(file.name, style="List Bullet")

    for task in spotlight_complete.tasks:
        document.add_heading(task.title, level=1)

        uuid_to_file_map = {f.uuid: f for f in files}

        raw = task.raw
        for uuid in uuid_to_file_map.keys():
            raw = raw.replace(f"<Doc{uuid}>", f"{uuid_to_file_map[uuid].name}")
            raw = raw.replace(f"</Doc{uuid}>", "")

            raw = raw.replace(f"Doc{uuid}", f"{uuid_to_file_map[uuid].name}")
            raw = raw.replace(f"{uuid}", f"{uuid_to_file_map[uuid].name}")

        html_raw = markdown.markdown(task.raw)
        temp_file = tempfile.NamedTemporaryFile(delete=True, suffix=".html")

        with open(temp_file.name, "w", encoding="utf-8") as f:
            f.write(html_raw)

        elements = partition_html(temp_file.name)

        for element in elements:
            element_dict = element.to_dict()
            if element_dict["type"] == "NarrativeText":
                document.add_paragraph(element_dict["text"])
            elif element_dict["type"] == "ListItem":
                indent = lookup_indentedness(
                    raw.replace("**", "").replace("__", ""),
                    line_str_to_match=element_dict["text"],
                )
                if indent is None:
                    indent = 0
                para = document.add_paragraph(element_dict["text"], style="List Bullet")
                para.paragraph_format.left_indent = Inches(indent / 4)
            elif element_dict["type"] == "Title":
                document.add_heading(element_dict["text"], level=2)
            else:
                print(element_dict["type"], element_dict["text"])

        document.add_page_break()

    return document
